from bs4 import BeautifulSoup, NavigableString
import os
import re
import requests
from PIL import Image
from io import BytesIO
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Inches
import time
import logging
from typing import Dict, List
from concurrent.futures import ThreadPoolExecutor
import urllib.parse
import httpx
import mammoth
from langdetect import detect
from html import unescape

def extract_navigation_paths(html):
    """
    Extract navigation paths from HTML content.
    This function finds all divs with class Step_1 and extracts navigation paths
    present in the HTML file content.
    
    Args:
        html (str): HTML content to parse
        
    Returns:
        list: List of navigation paths found in the HTML
    """
    soup = BeautifulSoup(html, "html.parser")
    nav_paths = []
 
    # Find all divs with class Step_1
    for div in soup.find_all("div", class_="Step_1"):
        parts = []
        for elem in div.children:
            if isinstance(elem, NavigableString):
                txt = elem
                print(f"we are printing txt for navigation path{txt}")
                if txt == ">" or txt == "&gt;" or txt==" > ":
                    if txt=="&gt;":
                        decoded_text = unescape(txt)
                        parts.append(decoded_text)
                    else:
                        parts.append(txt)
               
            else:
                if elem.get("class") and "Command_002c_menucascade_002c_uicontrol" in elem.get("class"):
                    parts.append(elem.get_text(strip=True))
        if parts:
            nav_paths.append("".join(parts).strip())
    return nav_paths

def check_note(directory,docx_output_file):
    full_flag=False
    docx_output_file.add_heading('Note Validation', level=1)
    docx_output_file.add_heading("Description", level=2)
    docx_output_file.add_paragraph("We are validating the HTML files to ensure that all instances of Notes are properly enclosed within designated grey boxes")
    for root, dirs, files in os.walk(directory):
        html_files_invalid_note=[]
        for filename in files:
            if filename.endswith(".html"):
                html_file = os.path.join(root, filename)
                
                flag=False
                with open(html_file, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f, 'html.parser')

                note_divs = soup.find_all('div', class_=re.compile(r'Note'))
                found_note_in_note_div = False
                found_note_outside_note_div = False

                # Check inside <div class="Note">
                for div in note_divs:
                    if "Note:" in div.get_text():
                        found_note_in_note_div = True

                # Check elsewhere in the document
                for element in soup.find_all(string=lambda text: "Note:" in text):
                    parent = element.find_parent()
                    if parent:
                        if not parent.find_parent('div', class_=re.compile(r'Note')) and parent.get('class') != ['Note']:
                            found_note_outside_note_div = True
                            # print(f"[WARNING] 'Note:' found outside grey box in: {html_file}")

                if found_note_outside_note_div:
                    full_flag=True
                    flag=True
                    # print(f"[INVALID] 'Note:' found outside grey box in: {html_file}")
                if flag:
                    html_files_invalid_note.add(f"{html_file}")
                    # print(f"{html_file}")
                    # print(f"[INVALID] 'Note:' found outside grey box")
    if not full_flag:
        docx_output_file.add_paragraph("All instances of Notes are properly enclosed within designated grey boxes")
    else:
        table = docx_output_file.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        # Add headers to the table
        headers = ["Serial Number", "html_file", "Status","Remarks"]

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ""# Clear existing text
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
        len=len(html_files_invalid_note)
        for i in range(len):
            row = table.add_row()
            row.cells[0].text = str(i+1)
            row.cells[1].text = html_files_invalid_note[i]
            note_status="❌   Invalid"

            row.cells[2].text = note_status
            if"❌" in note_status:
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            
            row.cells[3].text = "'Note:' found outside grey box"

def validate_images_in_folder(folder_path,docx_output_file):
    docx_output_file.add_heading('Images Validation', level=1)
    docx_output_file.add_heading("Description", level=2)
    docx_output_file.add_paragraph("We are currently validating the images embedded in HTML files to determine whether they are scrambled or corrupted.")
    docx_output_file.add_paragraph("Not Validated :- This script is currently unable to validate that particular image it should be reviewed manually.")
    flag=False
    html_invalid_images=[]
    html_images=[]
    reason_corrupt=[]
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(".html"):
                html_path = os.path.join(root, file)
                results = validate_images_in_html(html_path, root)
                
                for r in results:
                    if r['status']=="Invalid":
                        flag=True
                        html_invalid_images.append(f"{r['html_file']}")
                        html_images.append(f"{r['image_src']}")
                        reason_corrupt.append(f"{r['reason']}")               
                    
    if not flag:
        docx_output_file.add_paragraph("No image in any html file is scrambled or corrupted.")
    else:
        length = len(html_invalid_images)
        table = docx_output_file.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Add headers to the table
        headers = ["Serial Number", "html_file", "Image source","Status","Remarks"]

        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ""# Clear existing text
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
        
        for i in range(length):
            row = table.add_row()
            row.cells[0].text = str(i+1)
            # row.cells[1].text = html_invalid_images[i]
            row.cells[1].text = os.path.relpath(html_invalid_images[i], folder_path)
            row.cells[2].text = os.path.relpath(html_images[i], folder_path) if os.path.isabs(html_images[i]) else html_images[i]
            # row.cells[2].text = html_images[i]
            remarks1=reason_corrupt[i]
            if "Unreadable/corrupt image" in remarks1:
                image_status="⚠   Not Validated"
                
            else:
                image_status="❌   Invalid"
                

            row.cells[3].text = image_status
            if"⚠" in image_status:
                row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 165, 0)  #Orange
            elif "❌" in image_status:
                row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0) #Red
            row.cells[4].text = reason_corrupt[i]

def check_bullets(directory, docx_output_file):
    """Check bullet alignment in all HTML files in the directory"""
    docx_output_file.add_heading('Bullet Validation', level=1)
    docx_output_file.add_heading("Description", level=2)
    docx_output_file.add_paragraph("We are checking all the bullets inside HTML page, whether they are properly aligned or not.")
    
    files_checked = 0
    html_files_with_issues = []
    
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith('.html'):
                files_checked += 1
                file_path = os.path.join(root, file)
                
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        soup = BeautifulSoup(f, 'html.parser')
                        bullet_elements = set(
                            soup.find_all('div', class_=re.compile(r'List_\d+_-_\w+')) +
                            soup.find_all('li')
                        )
                        
                        if not bullet_elements:
                            continue
                            
                        misaligned = []
                        for element in bullet_elements:
                            style = element.get('style', '')
                            
                            # Check margin-left for List_ classes
                            if 'List_' in element.get('class', ''):
                                margin_match = re.search(r'margin-left:\s*(\d+(?:\.\d+)?)(pt|px)', style)
                                if not margin_match:
                                    misaligned.append((element, "Missing margin-left property"))
                                elif margin_match.group(1) != "18" or margin_match.group(2) != "pt":
                                    misaligned.append((element, f"Invalid margin-left: {margin_match.group(1)}{margin_match.group(2)}"))
                            
                            # Check for problematic properties
                            invalid_props = re.findall(r'(padding|line-height):\s*[^;]+', style)
                            if invalid_props:
                                misaligned.append((element, f"Found problematic properties: {', '.join(invalid_props)}"))
                        
                        if misaligned:
                            html_files_with_issues.append({
                                'file': file_path,
                                'total': len(bullet_elements),
                                'misaligned': misaligned
                            })
                            
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
                    continue
    
    if not html_files_with_issues:
        p = docx_output_file.add_paragraph()
        check_run = p.add_run("✓ ")
        check_run.font.color.rgb = RGBColor(0, 128, 0)
        p.add_run("All bullet points in all files are aligned properly.")
    else:
        for issue in html_files_with_issues:
            rel_path = os.path.relpath(issue['file'], directory)
            docx_output_file.add_heading(f"File: {rel_path}", level=3)
            
            summary = f"✗ Misaligned bullets found: {len(issue['misaligned'])} out of {issue['total']}"
            p = docx_output_file.add_paragraph(summary)
            p.style = 'Intense Quote'
            
            for idx, (element, reason) in enumerate(issue['misaligned'], 1):
                text = element.get_text(strip=True)
                docx_output_file.add_paragraph(f"{idx}. ID: {element.get('id', '')}")
                docx_output_file.add_paragraph(f"   Text: \"{text[:50]}...\"" if len(text) > 50 else f"   Text: \"{text}\"")
                
                reason_para = docx_output_file.add_paragraph("   Reason: ")
                reason_run = reason_para.add_run(reason)
                reason_run.font.color.rgb = RGBColor(255, 0, 0)
                
                docx_output_file.add_paragraph("_" * 40)

def is_menu_link(a_tag):
    """Check if the <a> tag is inside a menu-like structure."""
    menu_keywords = ['nav', 'menu', 'navbar', 'header', 'sidebar']
    for parent in a_tag.parents:
        if parent.name in ['nav', 'header']:
            return True
        if any(keyword in (parent.get('class') or []) for keyword in menu_keywords):
            return True
        if any(keyword in (parent.get('id') or '') for keyword in menu_keywords):
            return True
    return False
 
def is_skip_link(a_tag):
    """Check if the <a> tag is a skip link or similar."""
    if a_tag.get('aria-label') and 'Skip to Main content' in a_tag.get('aria-label'):
        return True
    return False
 
def check_spaces_in_html_folder(folder_path, docx_output_file):
    report = []
 
    for root, _, files in os.walk(folder_path):
        for filename in files:
            if filename.endswith('.html'):
                file_path = os.path.join(root, filename)
                try:
                    with open(file_path, 'r', encoding='utf-8') as file:
                        html_content = file.read()
 
                    soup = BeautifulSoup(html_content, 'html.parser')
 
                    for a_tag in soup.find_all('a'):
                        if is_menu_link(a_tag) or is_skip_link(a_tag):
                            continue  # Skip menu and skip links
 
                        parent_text = a_tag.parent.get_text()
                        link_text = a_tag.get_text()
 
                        match = re.search(re.escape(link_text), parent_text)
                        if match:
                            start, end = match.span()
                            before = parent_text[start - 1] if start > 0 else ''
                            after = parent_text[end] if end < len(parent_text) else ''
 
                            has_space_before = before.isspace()
                            has_space_after = after.isspace() if after != '.' else False
 
                            if has_space_before or has_space_after:
                                report.append({
                                    'file': os.path.relpath(file_path, folder_path),
                                    'link_text': link_text,
                                    'space_before': has_space_before,
                                    'space_after': has_space_after,
                                    'remarks': 'Missing space detected'
                                })
                except Exception as e:
                    print(f"Error reading {file_path}: {e}")
 
    # Write to Word document
    docx_output_file.add_heading('Spacing Detection Between Link and Text', level=1)
 
    if report:
        table = docx_output_file.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ["HTML File", "Link", "Missing Space Before", "Missing Space After", "Remarks"]
 
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = ""
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
 
        for entry in report:
            row = table.add_row()
            row.cells[0].text = entry['file']
            row.cells[1].text = entry['link_text']
            row.cells[2].text = str(entry['space_before'])
            row.cells[3].text = str(entry['space_after'])
            row.cells[4].text = entry['remarks']
    else:
        docx_output_file.add_paragraph("No missing spaces detected between link and text.")

def validate_images_in_html(html_path, base_path):
    results = []
    with open(html_path, 'r', encoding='utf-8') as f:
        soup = BeautifulSoup(f, 'html.parser')

    img_tags = soup.find_all('img')

    for img in img_tags:
        src = img.get('src')
        entry = {
            "html_file": html_path,
            "image_src": src if src else "MISSING",
            "status": "Valid",
            "reason": ""
        }
        if not src:
            entry["status"] = "Invalid"
            entry["reason"] = "Image tag without src attribute"
            results.append(entry)
            continue
        try:
            if src.startswith(('http://', 'https://')):
                response = requests.get(src, timeout=5)
                if response.status_code != 200:
                    entry["status"] = "Invalid"
                    entry["reason"] = f"Image URL not reachable (status {response.status_code})"
                else:
                    Image.open(BytesIO(response.content)).verify()
            else:
                image_path = os.path.normpath(os.path.join(base_path, src))
                if not os.path.exists(image_path):
                    entry["status"] = "Invalid"
                    entry["reason"] = f"Local image not found: {os.path.relpath(image_path, base_path)}"
                else:
                    with Image.open(image_path) as img_file:
                        img_file.verify()
        except Exception as e:
            entry["status"] = "Invalid"
            error_msg = str(e)
            if "image_path" in locals():
                rel_image_path = os.path.relpath(image_path, base_path)
                pattern = re.escape(os.path.basename(image_path)) + r"'"
                error_msg = re.sub(
                    r"[A-Za-z]:[\\/].*?" + pattern,
                    rel_image_path.replace("\\", "/") + "'",
                    error_msg)
                # error_msg = error_msg.replace(image_path, rel_image_path)
                # error_msg = error_msg.replace(image_path.replace('\\', '/'), rel_image_path)
                # error_msg = error_msg.replace(image_path.replace('/', '\\'), rel_image_path)
            entry["reason"] = f"Unreadable/corrupt image | Error: {error_msg}"
        results.append(entry)
    return results


def create_word_report(validation_results_by_file, output_file):
    """Create table report in Word document"""
    
    # Create main table with headers
    table = output_file.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Add headers
    headers = ["File Name", "Issue Category", "Issue Description", "Occurrences"]
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
    
    # Track current file for merging cells
    current_file = None
    
    for file_name, results in validation_results_by_file.items():
        if "error" in results or not any(results.get(k) for k in ["structure_issues", "boundary_issues", "image_issues"]):
            continue
            
        for category, issues_list in results.items():
            if not issues_list or category in ["alignment_issues", "formatting_issues"]:
                continue
                
            # Group similar issues and count occurrences
            issue_counts = {}
            for issue in issues_list:
                base_issue = re.sub(r'Table \d+, (Cell|Row|Column|Image) \d+:', '', issue)
                base_issue = re.sub(r'Table \d+:', '', base_issue).strip()
                issue_counts[base_issue] = issue_counts.get(base_issue, 0) + 1
            
            # Add rows to table
            for base_issue, count in issue_counts.items():
                row = table.add_row()
                
                # Only add file name if it's a new file
                if file_name != current_file:
                    row.cells[0].text = file_name
                    current_file = file_name
                
                row.cells[1].text = category.replace('_', ' ').title()
                row.cells[2].text = base_issue
                row.cells[3].text = str(count)


def link_checker(root_dir, output_file, threads=10):
    
    root_dir = os.path.abspath(root_dir)
    results = []
    checked_urls = set()  # To avoid checking duplicate URLs
    local_file_map = {}   # To map relative paths to absolute paths
    
    # Create a mapping of all HTML files in the directory structure
    for dirpath, _, filenames in os.walk(root_dir):
        for filename in filenames:
            if filename.lower().endswith('.html'):
                abs_path = os.path.join(dirpath, filename)
                rel_path = os.path.relpath(abs_path, root_dir)
                local_file_map[rel_path.replace('\\', '/')] = abs_path
                
                # Also map with a leading slash
                local_file_map['/' + rel_path.replace('\\', '/')] = abs_path

    # Find all HTML files in the directory structure
    html_files = []
    for dirpath, _, filenames in os.walk(root_dir):
        for filename in filenames:
            if filename.lower().endswith('.html'):
                html_files.append(os.path.join(dirpath, filename))
                
    # Extract links from HTML files
    all_links = []
    for html_file in html_files:
        try:
            with open(html_file, 'r', encoding='utf-8') as f:
                content = f.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            base_url = os.path.dirname(html_file)
            
            # Define extensions to exclude
            excluded_extensions = ('.css', '.js', '.svg', '.png', '.jpg', '.jpeg', '.gif', '.ico')
            
            links = []
            # Find all <a> tags with href
            for a_tag in soup.find_all('a', href=True):
                href = a_tag.get('href')
                if href and not href.startswith('javascript:') and not href.startswith('mailto:'):
                    # Skip links with excluded extensions and empty links
                    if not href.lower().endswith(excluded_extensions) and href != '#':
                        links.append({
                            'source_file': html_file,
                            'href': href,
                            'text': a_tag.get_text(strip=True)[:50],
                            'base_url': base_url
                        })
                        
            # Find all <link> tags with href (optional, can be removed if not needed)
            for link_tag in soup.find_all('link', href=True):
                href = link_tag.get('href')
                if href and not href.lower().endswith(excluded_extensions):
                    links.append({
                        'source_file': html_file,
                        'href': href,
                        'text': 'CSS/Resource link',
                        'base_url': base_url
                    })
                        
            # Find all <img> tags with src (optional, can be removed if not needed)
            for img_tag in soup.find_all('img', src=True):
                src = img_tag.get('src')
                if src and not src.lower().endswith(excluded_extensions):
                    links.append({
                        'source_file': html_file,
                        'href': src,
                        'text': img_tag.get('alt', 'Image')[:50],
                        'base_url': base_url
                    })
                    
            # Remove duplicates links
            unique_links = []    
            seen = set()
            for link in links:
                if link['href'] not in seen:
                    seen.add(link['href'])
                    unique_links.append(link)

            all_links.extend(unique_links)
        
        except Exception as e:
            print(f"Error extracting links from {html_file}: {e}")
    
    # Helper function to resolve a URL to its absolute form
    def resolve_url(href, base_url):
        if href.startswith(('http://', 'https://')):
            return href
        
        # Special case for fragment-only URLs
        if href.startswith('#'):
            return os.path.relpath(base_url, root_dir) + href
                
        # For relative URLs, resolve against base_url
        if not os.path.isabs(href):
            if href.startswith('/'):
                # Absolute path relative to root directory
                resolved = os.path.join(root_dir, href[1:])
            else:
                # Relative path
                resolved = os.path.normpath(os.path.join(base_url, href))
        else:
            resolved = href
                
        # Return path relative to root directory for consistency
        return os.path.relpath(resolved, root_dir).replace('\\', '/')
    
    # Helper function to check if a link is valid
    def check_link(link_info):
        href = link_info['href']
        source_file = link_info['source_file']
        base_url = link_info['base_url']
        
        # Skip checking if we've already checked this URL
        resolved_url = resolve_url(href, base_url)
        if resolved_url in checked_urls:
            # Get the status from a previous check
            for result in results:
                if result['resolved_url'] == resolved_url:
                    return {
                        'source_file': os.path.relpath(source_file, root_dir),
                        'href': href,
                        'text': link_info['text'],
                        'resolved_url': resolved_url,
                        'status': result['status'],
                        'is_broken': result['is_broken'],
                        'error': result.get('error', '')
                    }        
        checked_urls.add(resolved_url)
        
        result = {
            'source_file': os.path.relpath(source_file, root_dir),
            'href': href,
            'text': link_info['text'],
            'resolved_url': resolved_url
        }
        
        try:
            if href.startswith(('http://', 'https://')):
                # External URL
                response = requests.head(href, allow_redirects=True, timeout=5)
                if response.status_code >= 400:
                    response = requests.get(href, timeout=5)  # Try GET if HEAD fails
                result['status'] = response.status_code
                result['is_broken'] = response.status_code >= 400
                
            else:
                # Local file or relative URL
                is_broken = True
                error_message = "File not found"
                
                # Handle fragment identifiers like index.html#page/section.html
                parts = href.split('#')
                base_href = parts[0]
                
                if not base_href:  # Just a fragment like "#section"
                    # This is likely a link within the same page
                    is_broken = False
                    error_message = ""
                else:
                    # Handle relative paths
                    if not os.path.isabs(base_href):
                        if base_href.startswith('/'):
                            # Absolute path relative to root
                            local_path = os.path.join(root_dir, base_href[1:])
                        else:
                            # Relative path
                            local_path = os.path.normpath(os.path.join(base_url, base_href))
                    else:
                        local_path = base_href
                    
                    # Convert to proper path format
                    local_path = local_path.replace('\\', '/')
                    
                    # Handle special case with encoded URLs
                    decoded_path = urllib.parse.unquote(local_path)
                    
                    if os.path.exists(local_path):
                        is_broken = False
                        error_message = ""
                    elif os.path.exists(decoded_path):
                        is_broken = False
                        error_message = ""
                    else:
                        # Check if the path exists in our file map
                        rel_path = os.path.relpath(local_path, root_dir).replace('\\', '/')
                        if rel_path in local_file_map or '/'+rel_path in local_file_map:
                            is_broken = False
                            error_message = ""
                
                result['status'] = 200 if not is_broken else 404
                result['is_broken'] = is_broken
                if is_broken:
                    result['error'] = error_message
                
        except requests.exceptions.Timeout:
            result['status'] = 'Timeout'
            result['is_broken'] = True
            result['error'] = 'Request timed out'
        except requests.exceptions.ConnectionError:
            result['status'] = 'Connection Error'
            result['is_broken'] = True
            result['error'] = 'Failed to establish connection'
        except Exception as e:
            result['status'] = 'Error'
            result['is_broken'] = True
            result['error'] = str(e)
            
        return result
    
    # Check links in parallel
    with ThreadPoolExecutor(max_workers=threads) as executor:
        results = list(executor.map(check_link, all_links))
    
    # Get status text function
    def get_status_text(status_code):
        """Get the text description for an HTTP status code"""
        status_descriptions = {
            400: "Bad Request",
            401: "Unauthorized",
            403: "Forbidden",
            404: "Not Found",
            500: "Internal Server Error",
            502: "Bad Gateway",
            503: "Service Unavailable",
            504: "Gateway Timeout"
        }
        return status_descriptions.get(status_code, f"HTTP Error {status_code}")

    # Write report to Word document
    output_file.add_heading('Link Validation', level=1)
    output_file.add_heading("Description", level=2)
    output_file.add_paragraph("We are validating all links in HTML files to ensure they are accessible and properly formatted.")
    output_file.add_paragraph("**Not Validated: Access is restricted for that URL.**")

    broken_links = [r for r in results if r['is_broken']]
    
    if broken_links:
        # Group broken links by URL and link text
        consolidated_links = {}
        for result in broken_links:
            # Create a key based on link text, URL, and status
            key = (result['text'], result['href'], result['status'])
            
            # Add to consolidated dictionary
            if key in consolidated_links:
                consolidated_links[key]['source_files'].append(result['source_file'])
            else:
                consolidated_links[key] = {
                    'text': result['text'],
                    'href': result['href'],
                    'status': result['status'],
                    'source_files': [result['source_file']],
                    'error': result.get('error', '')
                }
        
        # Create table with consolidated data
        # table = doc.add_table(rows=1, cols=5, style='Table Grid')
        table = output_file.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set column widths (as percentage of page width)
        col_widths = [15, 20, 10, 15, 40]  # percentages
        for i, width in enumerate(col_widths):
            table.columns[i].width = Inches(width * 6.0 / 100)  # 6.0 inches is typical usable page width
            
        # Add headers
        headers = ["Link Name", "URL", "Status", "Remarks", "HTML Page"]
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.paragraphs[0].runs[0].bold = True
        
        # Add data rows
        for key, data in consolidated_links.items():
            row = table.add_row().cells
            
            # Link name
            row[0].text = data['text']
            
            # URL
            row[1].text = data['href']
            
            # Status with emoji indicator
            status_code = data['status']
            status_cell = row[2]
            status_para = status_cell.paragraphs[0]
            
            # Handle different status codes
            if isinstance(status_code, str):
                status_text, emoji, color = "Not Validated", "⚠", RGBColor(255, 165, 0)
                remark = f"HTTP {status_code}: {data.get('error', 'Unknown error')}"
            else:
                if status_code == 403:
                    status_text, emoji, color = "Not Validated", "⚠", RGBColor(255, 165, 0)
                    remark = f"HTTP {status_code}: Forbidden"
                elif status_code >= 400:
                    status_text, emoji, color = "Not Validated", "⚠", RGBColor(255, 165, 0)
                    remark = f"HTTP {status_code}: {get_status_text(status_code)}"
                else:
                    status_text, emoji, color = "Correct", "✓", RGBColor(0, 128, 0)
                    remark = f"HTTP {status_code}: OK"
            
            # Add status with emoji and coloring
            run = status_para.add_run(f"{emoji}\nNot Validated")
            run.font.color.rgb = color
            status_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Remarks
            row[3].text = remark
            
            # HTML Pages (comma-separated list)
            html_pages_cell = row[4]
            for i, file_path in enumerate(data['source_files']):
                if i > 0:
                    html_pages_cell.paragraphs[0].add_run(",\n")
                html_pages_cell.paragraphs[0].add_run(file_path)

    else:
        output_file.add_paragraph("No broken links found.")

    # Save the document
    return results


def is_camel_case(s):
    if not isinstance(s, str):
        return False
 
    # Remove brackets from the string
    brackets = "[]{}()"
    s_clean = ''.join(c for c in s if c not in brackets)
 
    return (
        s_clean != s_clean.lower() and
        s_clean != s_clean.upper() and
        "_" not in s_clean and
        " " not in s_clean and
        any(c.isupper() for c in s_clean[1:])
    )

def check_navigation_path(folder_path, docx_output_file):
    """
    Validate navigation paths in HTML files.
    Checks for:
    1. Correct format with " > " (single space before and after >)
    2. No camel case words in navigation paths
    """
    docx_output_file.add_heading('Navigation Path Validation', level=1)
    docx_output_file.add_heading("Description", level=2)
    docx_output_file.add_paragraph("We are validating the HTML files to ensure that all navigation paths are written properly with correct spacing and no camel case words")
 
    issues = []
 
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(".html"):
                html_path = os.path.join(root, file)
                file_name = os.path.basename(html_path)
 
                try:
                    with open(html_path, 'r', encoding='utf-8') as f:
                        html_content = f.read()
                except Exception as e:
                    print(f"Error reading {html_path}: {e}")
                    continue
 
                # Extract navigation paths using our custom function
                nav_paths = extract_navigation_paths(html_content)
               
                for nav_path in nav_paths:
                    # Check for correct spacing around ">"
                    # Should have " > " (space before and after)
                    incorrect_spacing = False
                    if ">" in nav_path:
                        # Check for patterns like "text>text" or "text >text" or "text> text"
                        if re.search(r'\S>\S|[^\s]>\s|\s>[^\s]', nav_path):
                            incorrect_spacing = True
                   
                    # Check for camel case words
                    # Split by " > " to get individual parts and check each part
                    parts = re.split(r'\s*>\s*', nav_path)
                    has_camel_case = False
                    camel_case_words = []
                   
                    for part in parts:
                        words = part.split()
                        for word in words:
                           
                            sub_words = re.split(r"[\/\\\-_'\(\)]", word)
                            for sub_word in sub_words:
                                if sum(1 for c in sub_word if c.isupper()) >= 2:
 
                                    camel_case_words.append(word)
                   
                    # Add issues to the report
                    if incorrect_spacing:
                        issues.append({
                            "file": os.path.relpath(html_path, folder_path),
                            "line": "-",
                            "issue": nav_path,
                            "status": "❌   Invalid",
                            "remarks": "Incorrect spacing around '>' - should be ' > ' (single space before and after)"
                        })
                   
                    if has_camel_case:
                        issues.append({
                            "file": os.path.relpath(html_path, folder_path),
                            "line": "-",
                            "issue": nav_path,
                            "status": "❌   Invalid",
                            "remarks": f"Contains words with multiple capital letters: {', '.join(camel_case_words)}"
                        })
 
    # Create table
    if issues:
        table = docx_output_file.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ["Serial Number", "HTML File", "Navigation Path", "Status", "Remarks"]
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
 
        for i, issue in enumerate(issues):
            row = table.add_row()
            row.cells[0].text = str(i + 1)
            row.cells[1].text = issue["file"]
            row.cells[2].text = issue["issue"]
            row.cells[3].text = issue["status"]
            if "❌" in issue["status"]:
                row.cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
            row.cells[4].text = issue["remarks"]
    else:
        docx_output_file.add_paragraph("✅ No issues found in navigation paths.")

def extract_docx_html(docx_path):
    """Extracts the content of a DOCX file and converts it to HTML."""
    with open(docx_path, "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value
    return html

def run_usecase(usecase_name, form_data, exec_id, job_name, model_name, input_dirpath, output_dirpath, *args, **kwargs):   
    print("outputFolder............."+output_dirpath)
    os.makedirs(output_dirpath, exist_ok=True)
    # Ensure the output directory exists

    t1 = time.time()
    # Create a new Word document for each PDF file
    docx_output_file = Document()
    heading = docx_output_file.add_heading(level=1)
    run = heading.add_run(f"Spot Check validation for html files")
    run.bold = True
    
    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center align the heading
        
    # Perform checks and generate the report
    validate_images_in_folder(input_dirpath,docx_output_file)
    check_spaces_in_html_folder(input_dirpath, docx_output_file)
    link_checker(input_dirpath, docx_output_file)
    check_note(input_dirpath,docx_output_file)
    check_bullets(input_dirpath, docx_output_file)
    check_navigation_path(input_dirpath,docx_output_file)
    input_folder_name = os.path.basename(output_dirpath.rstrip(os.sep))
    report_name = f"{input_folder_name}_Report.docx"

    docx_path=output_dirpath
    docx_output_file_path = os.path.join(docx_path, report_name)
    docx_output_file.save(docx_output_file_path)

    t2 = time.time()
    timetaken = t2 - t1

    report_html = extract_docx_html(docx_output_file_path)
 
    num_of_files = sum(len(files) for _, _, files in os.walk(input_dirpath))
 
    return {"status": "success", "msg": "Success"}

if __name__ == "__main__":
    run_usecase(usecase_name="",form_data="", exec_id="", job_name="", model_name="",input_dirpath="C:/Users/pasiratnesh.tarakant/Downloads/Xerox_en-US_initial_HTML/Xerox_en-US", output_dirpath="FINAL_HTML_OUTPUT")